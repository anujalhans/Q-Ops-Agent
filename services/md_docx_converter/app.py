from flask import Flask, request, send_file, jsonify, after_this_request
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from markdown import markdown
from bs4 import BeautifulSoup
import tempfile
import os
import datetime
import re
import time
import logging
import uuid
from io import BytesIO

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 2 * 1024 * 1024

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
LOGO_PATH = os.environ.get(
    "CONVERTER_LOGO_PATH",
    os.path.join(BASE_DIR, "assets", "royal_enfield_logo.png")
)
MAX_MARKDOWN_CHARS = 1_000_000
MAX_CONVERSION_SECONDS = 120
WIDE_CONFLUENCE_TABLE_COLUMN_THRESHOLD = 6
WIDE_CONFLUENCE_TABLE_MIN_WIDTH_PX = 1400
WIDE_CONFLUENCE_TABLE_MAX_WIDTH_PX = 2200

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(levelname)s %(name)s %(message)s"
)
logger = logging.getLogger("md-to-docx-converter")


def log_event(event, request_id=None, **fields):
    details = " ".join(f"{key}={value}" for key, value in fields.items() if value is not None)
    prefix = f"request_id={request_id} " if request_id else ""
    logger.info("%s%s %s", prefix, event, details)


def check_duration(stage, started_at, request_id):
    elapsed = time.perf_counter() - started_at
    if elapsed > MAX_CONVERSION_SECONDS:
        raise TimeoutError(f"{stage} exceeded {MAX_CONVERSION_SECONDS}s")
    log_event("stage_completed", request_id, stage=stage, duration=f"{elapsed:.3f}s")
    return elapsed


def get_table_column_count(table):
    rows = table.find_all("tr")
    if not rows:
        return 0
    return max(len(row.find_all(["th", "td"])) for row in rows)


def is_wide_confluence_table(table):
    return get_table_column_count(table) >= WIDE_CONFLUENCE_TABLE_COLUMN_THRESHOLD


def apply_confluence_table_layout(soup, table):
    column_count = get_table_column_count(table)
    table['class'] = 'confluenceTable'

    if is_wide_confluence_table(table):
        min_width = min(
            WIDE_CONFLUENCE_TABLE_MAX_WIDTH_PX,
            max(WIDE_CONFLUENCE_TABLE_MIN_WIDTH_PX, column_count * 180)
        )
        table['data-layout'] = 'wide'
        table['style'] = f"min-width:{min_width}px; table-layout:auto;"

        wrapper = soup.new_tag("div")
        wrapper['class'] = 'qops-wide-table'
        wrapper['style'] = 'overflow-x:auto; width:100%;'
        table.wrap(wrapper)
    else:
        table['data-layout'] = 'default'


def register_temp_file_cleanup(path, request_id):
    @after_this_request
    def cleanup_temp_file(response):
        def remove_temp_file():
            try:
                if path and os.path.exists(path):
                    os.remove(path)
                    log_event("temp_file_deleted", request_id, path=path)
            except Exception as cleanup_error:
                logger.warning(
                    "request_id=%s temp_file_cleanup_failed path=%s error=%s",
                    request_id,
                    path,
                    cleanup_error
                )

        response.call_on_close(remove_temp_file)
        return response

    return cleanup_temp_file


def delete_temp_file(path, request_id):
    if path and os.path.exists(path):
        os.remove(path)
        log_event("temp_file_deleted", request_id, path=path)


def send_docx_response(path, file_name, request_id):
    with open(path, "rb") as docx_file:
        docx_bytes = docx_file.read()

    delete_temp_file(path, request_id)
    return send_file(
        BytesIO(docx_bytes),
        as_attachment=True,
        download_name=file_name,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


# ---------------------------------------------------
# STYLE CONFIGURATION
# ---------------------------------------------------

def configure_styles(doc):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 51, 102)

    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(14)
    h2.font.bold = True


# ---------------------------------------------------
# COVER PAGE (Dynamic)
# ---------------------------------------------------

def add_cover_page(doc, document_type):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = title.add_run(f"{document_type}\n")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0, 51, 102)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Project - ShopSmart")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()


# ---------------------------------------------------
# AUTO TABLE OF CONTENTS
# ---------------------------------------------------

def add_table_of_contents(doc):
    heading = doc.add_heading("Table of Contents", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)

    instrText = OxmlElement('w:instrText')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instrText)

    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)

    doc.add_page_break()


def force_update_fields(doc):
    settings = doc.settings
    update_fields = OxmlElement('w:updateFields')
    update_fields.set(qn('w:val'), 'true')
    settings._element.append(update_fields)


# ---------------------------------------------------
# HEADER / FOOTER
# ---------------------------------------------------

def add_logo_to_header(section):
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if os.path.exists(LOGO_PATH):
        run = paragraph.add_run()
        run.add_picture(LOGO_PATH, width=Inches(1.2))


def add_footer(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    paragraph.add_run(
        f"Generated on {datetime.datetime.now().strftime('%d-%b-%Y %H:%M')}  |  Confidential  |  Page "
    )

    run = paragraph.add_run()

    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)

    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"
    run._r.append(instrText)

    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)


# ---------------------------------------------------
# MARKDOWN → DOCX
# ---------------------------------------------------

def add_markdown_content(doc, md_text):
    html = markdown(md_text, extensions=['tables'])
    soup = BeautifulSoup(html, 'html.parser')

    for element in soup.children:
        if element.name == 'h1':
            doc.add_heading(element.get_text(), level=1)

        elif element.name == 'h2':
            doc.add_heading(element.get_text(), level=2)

        elif element.name == 'h3':
            doc.add_heading(element.get_text(), level=3)

        elif element.name == 'p':
            para = doc.add_paragraph(element.get_text())
            para.paragraph_format.space_after = Pt(8)

        elif element.name == 'ul':
            for li in element.find_all('li'):
                doc.add_paragraph(li.get_text(), style='List Bullet')

        elif element.name == 'ol':
            for li in element.find_all('li'):
                doc.add_paragraph(li.get_text(), style='List Number')

        elif element.name == 'table':
            rows = element.find_all('tr')
            if not rows:
                continue

            col_count = max(len(row.find_all(['th', 'td'])) for row in rows)
            if col_count == 0:
                continue

            table = doc.add_table(rows=len(rows), cols=col_count)
            table.style = 'Table Grid'

            for i, row in enumerate(rows):
                cells = row.find_all(['th', 'td'])
                for j, cell in enumerate(cells):
                    table.rows[i].cells[j].text = cell.get_text(strip=True)

                    if i == 0:
                        for p in table.rows[i].cells[j].paragraphs:
                            for run in p.runs:
                                run.bold = True

            doc.add_paragraph()


# ===================================================
# 🆕 MARKDOWN → CONFLUENCE STORAGE FORMAT
# ===================================================

def convert_markdown_to_confluence(md_text):
    html = markdown(md_text, extensions=['tables', 'fenced_code'])
    soup = BeautifulSoup(html, 'html.parser')

    for tag in soup.find_all(['h1', 'h2', 'h3']):
        tag['style'] = "color:#172B4D;"

    for table in soup.find_all("table"):
        apply_confluence_table_layout(soup, table)

        for th in table.find_all("th"):
            th['class'] = 'confluenceTh'

        for td in table.find_all("td"):
            td['class'] = 'confluenceTd'

    for ul in soup.find_all("ul"):
        ul['class'] = 'confluenceList'

    for ol in soup.find_all("ol"):
        ol['class'] = 'confluenceList'

    for pre in soup.find_all("pre"):
        code = pre.find("code")
        if code:
            macro = soup.new_tag("ac:structured-macro")
            macro.attrs["ac:name"] = "code"

            plain_text = soup.new_tag("ac:plain-text-body")
            safe_code_text = code.get_text().replace("]]>", "]]]]><![CDATA[>")
            plain_text.string = f"<![CDATA[{safe_code_text}]]>"

            macro.append(plain_text)
            pre.replace_with(macro)

    for blockquote in soup.find_all("blockquote"):
        panel = soup.new_tag("ac:structured-macro")
        panel.attrs["ac:name"] = "info"

        body = soup.new_tag("ac:rich-text-body")
        for child in list(blockquote.contents):
            body.append(child.extract())

        panel.append(body)
        blockquote.replace_with(panel)

    return f"<div>{str(soup)}</div>"


# ---------------------------------------------------
# API ENDPOINT
# ---------------------------------------------------

@app.route('/convert', methods=['POST'])
def convert_markdown_to_docx():
    temp_file_path = None
    request_id = request.headers.get("X-Request-ID") or str(uuid.uuid4())
    try:
        request_started_at = time.perf_counter()
        if request.content_length and request.content_length > app.config["MAX_CONTENT_LENGTH"]:
            return jsonify({
                "error": "Request payload too large",
                "requestId": request_id,
                "maxBytes": app.config["MAX_CONTENT_LENGTH"]
            }), 413

        data = request.get_json(silent=True) or {}

        # ✅ SUPPORT BOTH FORMATS (IMPORTANT FIX)
        md_text = data.get('cleanedMarkdown') or data.get('markdown', '')
        document_type = data.get('documentType') or data.get('body', {}).get('documentType', 'Enterprise Document')
        response_type = data.get("responseType", "docx")

        if not md_text:
            return jsonify({"error": "No markdown content provided", "requestId": request_id}), 400

        if len(md_text) > MAX_MARKDOWN_CHARS:
            return jsonify({
                "error": "Markdown content too large",
                "requestId": request_id,
                "maxMarkdownChars": MAX_MARKDOWN_CHARS
            }), 413

        log_event(
            "request_started",
            request_id,
            documentType=document_type,
            responseType=response_type,
            chars=len(md_text)
        )

        doc = Document()
        configure_styles(doc)
        add_cover_page(doc, document_type)
        add_table_of_contents(doc)
        stage_started_at = time.perf_counter()
        add_markdown_content(doc, md_text)
        check_duration("add_markdown_content", stage_started_at, request_id)

        for section in doc.sections:
            add_logo_to_header(section)
            add_footer(section)

        force_update_fields(doc)

        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        temp_file_path = temp_file.name
        temp_file.close()

        stage_started_at = time.perf_counter()
        doc.save(temp_file_path)
        check_duration("doc.save", stage_started_at, request_id)

        safe_name = re.sub(r'[^A-Za-z0-9_]+', '_', document_type)
        file_name = f"{safe_name}.docx"

        # ✅ NEW LOGIC
        if response_type == "docx":
            log_event(
                "request_completed",
                request_id,
                responseType=response_type,
                duration=f"{time.perf_counter() - request_started_at:.3f}s"
            )
            response = send_docx_response(temp_file_path, file_name, request_id)
            temp_file_path = None
            return response

        elif response_type == "json":
            stage_started_at = time.perf_counter()
            confluence_content = convert_markdown_to_confluence(md_text)
            check_duration("convert_markdown_to_confluence", stage_started_at, request_id)
            delete_temp_file(temp_file_path, request_id)
            temp_file_path = None
            log_event(
                "request_completed",
                request_id,
                responseType=response_type,
                duration=f"{time.perf_counter() - request_started_at:.3f}s"
            )
            return jsonify({
                # 🔥 PRESERVE ORIGINAL STRUCTURE
                "headers": data.get("headers"),
                "params": data.get("params"),
                "query": data.get("query"),
                "body": data.get("body"),
                "webhookUrl": data.get("webhookUrl"),
                "executionMode": data.get("executionMode"),
                "cleanedMarkdown": data.get("cleanedMarkdown"),

                # 🔥 NEW OUTPUT
                "fileName": file_name,
                "confluenceContent": confluence_content
            })

        response = send_docx_response(temp_file_path, file_name, request_id)
        temp_file_path = None
        return response

    except Exception as e:
        if temp_file_path:
            try:
                delete_temp_file(temp_file_path, request_id)
            except Exception:
                logger.exception("request_id=%s temp_cleanup_after_error_failed", request_id)
        logger.exception("request_id=%s conversion_failed", request_id)
        status_code = 504 if isinstance(e, TimeoutError) else 500
        return jsonify({"error": str(e), "requestId": request_id}), status_code

@app.route('/health', methods=['GET'])
def health():
    return jsonify({
        "status": "up & running",
        "service": "md-to-docx-converter"
    })


@app.route('/ready', methods=['GET'])
def ready():
    request_id = request.headers.get("X-Request-ID") or str(uuid.uuid4())
    started_at = time.perf_counter()
    try:
        sample_markdown = "# Converter Ready\n\n> Diagnostic quote\n\n| A | B |\n|---|---|\n| 1 | 2 |"
        doc = Document()
        configure_styles(doc)
        add_markdown_content(doc, sample_markdown)
        confluence_content = convert_markdown_to_confluence(sample_markdown)
        duration = time.perf_counter() - started_at
        log_event("ready_check_completed", request_id, duration=f"{duration:.3f}s")
        return jsonify({
            "status": "ready",
            "service": "md-to-docx-converter",
            "requestId": request_id,
            "durationSeconds": round(duration, 3),
            "docxParagraphs": len(doc.paragraphs),
            "docxTables": len(doc.tables),
            "confluenceInfoMacro": 'ac:name="info"' in confluence_content,
            "confluenceTable": "confluenceTable" in confluence_content
        })
    except Exception as e:
        logger.exception("request_id=%s ready_check_failed", request_id)
        return jsonify({
            "status": "not_ready",
            "service": "md-to-docx-converter",
            "requestId": request_id,
            "error": str(e)
        }), 500


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5050, debug=False)
