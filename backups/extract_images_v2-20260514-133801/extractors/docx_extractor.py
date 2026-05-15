from __future__ import annotations

import base64
import io
from typing import Any
import xml.etree.ElementTree as ET
import zipfile

from docx import Document

from extract_images_v2.extractors.table_utils import table_rows_to_markdown, table_rows_to_plain_text
from extract_images_v2.extractors.visual_detection import (
    annotate_visual_candidates,
    convert_office_bytes_to_pdf_bytes,
    limit_visual_candidates,
    render_pdf_bytes,
    resolve_vision_config,
)


WORD_NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
}


def _extract_docx_comments(zip_file: zipfile.ZipFile, *, file_name: str) -> list[dict[str, Any]]:
    if "word/comments.xml" not in zip_file.namelist():
        return []

    root = ET.fromstring(zip_file.read("word/comments.xml"))
    comments: list[dict[str, Any]] = []
    for comment_index, comment in enumerate(root.findall(".//w:comment", WORD_NS)):
        text = " ".join(
            node.text.strip()
            for node in comment.findall(".//w:t", WORD_NS)
            if node.text and node.text.strip()
        )
        author = comment.attrib.get(f"{{{WORD_NS['w']}}}author", "")
        if not text:
            continue
        comments.append(
            {
                "annotationId": f"{file_name}_comment_{comment_index}",
                "fileName": file_name,
                "pageNumber": None,
                "type": "docx-comment",
                "title": author,
                "subject": "comment",
                "content": text,
                "source": "docx-comment",
            }
        )
    return comments


def _extract_docx_notes(zip_file: zipfile.ZipFile, *, file_name: str, note_type: str) -> list[dict[str, Any]]:
    xml_name = f"word/{note_type}.xml"
    if xml_name not in zip_file.namelist():
        return []

    root = ET.fromstring(zip_file.read(xml_name))
    notes: list[dict[str, Any]] = []
    for note_index, note in enumerate(root):
        text = " ".join(
            node.text.strip()
            for node in note.findall(".//w:t", WORD_NS)
            if node.text and node.text.strip()
        )
        if not text:
            continue
        notes.append(
            {
                "annotationId": f"{file_name}_{note_type}_{note_index}",
                "fileName": file_name,
                "pageNumber": None,
                "type": note_type,
                "title": note_type,
                "subject": note_type,
                "content": text,
                "source": note_type,
            }
        )
    return notes


def _extract_docx_links(document: Document, *, file_name: str) -> list[dict[str, Any]]:
    links: list[dict[str, Any]] = []
    for link_index, rel in enumerate(document.part.rels.values()):
        if "hyperlink" not in rel.reltype:
            continue
        links.append(
            {
                "linkId": f"{file_name}_link_{link_index}",
                "fileName": file_name,
                "pageNumber": None,
                "uri": rel.target_ref,
                "targetPage": None,
                "kind": "external",
                "source": "docx-link",
            }
        )
    return links


def _extract_docx_tables(document: Document, *, file_name: str) -> list[dict[str, Any]]:
    tables: list[dict[str, Any]] = []
    for table_index, table in enumerate(document.tables):
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        markdown = table_rows_to_markdown(rows)
        plain_text = table_rows_to_plain_text(rows)
        if not markdown and not plain_text:
            continue
        tables.append(
            {
                "tableId": f"{file_name}_table_{table_index}",
                "fileName": file_name,
                "pageNumber": None,
                "tableIndex": table_index,
                "rowCount": len(rows),
                "columnCount": max((len(row) for row in rows), default=0),
                "markdown": markdown,
                "text": plain_text,
                "source": "docx-table",
            }
        )
    return tables


def _extract_docx_images(document: Document, *, file_name: str, doc_type: str) -> list[dict[str, Any]]:
    images: list[dict[str, Any]] = []
    image_index = 0
    for rel in document.part.rels.values():
        if "image" not in rel.target_ref:
            continue
        image_bytes = rel.target_part.blob
        extension = rel.target_ref.rsplit(".", 1)[-1].lower()
        images.append(
            {
                "fileName": f"{file_name}_image_{image_index}.{extension}",
                "imageId": f"{file_name}_img_{image_index}",
                "base64": base64.b64encode(image_bytes).decode(),
                "docType": f"{doc_type}-IMAGE" if doc_type != "UNKNOWN" else "IMAGE",
                "imageSource": "embedded-image",
                "pageNumber": None,
                "visualReason": ["embedded_image"],
                "mimeType": f"image/{extension}",
            }
        )
        image_index += 1
    return images


def _detect_docx_visual_indicators(zip_file: zipfile.ZipFile) -> list[str]:
    names = set(zip_file.namelist())
    reasons: list[str] = []

    if any(name.startswith("word/diagrams/") for name in names):
        reasons.append("smartart_or_diagram")
    if any(name.endswith((".emf", ".wmf", ".svg")) for name in names):
        reasons.append("vector_media")
    if any(name.startswith("word/drawings/") for name in names):
        reasons.append("drawingml_shapes")
    if any(name.startswith("word/activeX/") or name.startswith("word/controls/") for name in names):
        reasons.append("form_controls")

    return reasons


def extract_docx_document(*, content: bytes, context: dict[str, Any]) -> dict[str, Any]:
    file_name = context["fileName"]
    doc_type = context["docType"]
    vision_config = resolve_vision_config(context)

    document = Document(io.BytesIO(content))
    raw_text = "\n".join(
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text and paragraph.text.strip()
    )
    tables = _extract_docx_tables(document, file_name=file_name)
    detected_images = _extract_docx_images(document, file_name=file_name, doc_type=doc_type)
    images = annotate_visual_candidates(detected_images)
    links = _extract_docx_links(document, file_name=file_name)

    rendered_pages: list[dict[str, Any]] = []
    warnings: list[str] = []
    deferred_embedded_images = 0
    deferred_rendered_pages = 0

    with zipfile.ZipFile(io.BytesIO(content)) as zip_file:
        annotations = _extract_docx_comments(zip_file, file_name=file_name)
        annotations.extend(_extract_docx_notes(zip_file, file_name=file_name, note_type="footnotes"))
        annotations.extend(_extract_docx_notes(zip_file, file_name=file_name, note_type="endnotes"))

        visual_indicators = _detect_docx_visual_indicators(zip_file)
        if visual_indicators:
            pdf_bytes, warning = convert_office_bytes_to_pdf_bytes(content, suffix=".docx")
            if warning:
                warnings.append(warning)
            if pdf_bytes:
                rendered_pages = render_pdf_bytes(
                    pdf_bytes,
                    file_name=file_name,
                    doc_type=doc_type,
                    visual_reason=visual_indicators,
                    prefix="docx_rendered_page",
                    dpi=vision_config["renderDpi"],
                )

    images, deferred_embedded_images = limit_visual_candidates(
        images,
        max_items=vision_config["maxEmbeddedImagesPerDocument"],
        overflow_label=f"{file_name} embedded images",
        warnings=warnings,
    )
    rendered_pages, deferred_rendered_pages = limit_visual_candidates(
        rendered_pages,
        max_items=vision_config["maxRenderedPagesPerDocument"],
        overflow_label=f"{file_name} rendered pages",
        warnings=warnings,
    )

    return {
        "fileType": "docx",
        "pageCount": len(rendered_pages),
        "rawText": raw_text,
        "images": images,
        "renderedPages": rendered_pages,
        "tables": tables,
        "annotations": annotations,
        "links": links,
        "warnings": warnings,
        "extractionStats": {
            "paragraphsExtracted": sum(1 for paragraph in document.paragraphs if paragraph.text and paragraph.text.strip()),
            "embeddedImagesDetected": len(detected_images),
            "embeddedImagesExtracted": len(images),
            "embeddedImagesDeferred": deferred_embedded_images,
            "tablesExtracted": len(tables),
            "renderedPageCandidatesDetected": len(rendered_pages) + deferred_rendered_pages,
            "renderedPagesGenerated": len(rendered_pages),
            "renderedPagesDeferred": deferred_rendered_pages,
            "annotationsExtracted": len(annotations),
            "linksExtracted": len(links),
            "visualCandidatesDetected": len(images) + deferred_embedded_images + len(rendered_pages) + deferred_rendered_pages,
            "visualCandidatesReturned": len(images) + len(rendered_pages),
            "visualCandidatesDeferred": deferred_embedded_images + deferred_rendered_pages,
            "renderDpi": vision_config["renderDpi"],
        },
    }
