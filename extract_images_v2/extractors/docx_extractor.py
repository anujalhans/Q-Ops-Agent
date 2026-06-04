from __future__ import annotations

import base64
import io
from typing import Any
import xml.etree.ElementTree as ET
import zipfile

from docx import Document

from extract_images_v2.extractors.table_utils import table_rows_to_markdown, table_rows_to_plain_text
from extract_images_v2.extractors.visual_detection import annotate_visual_candidate


WORD_NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
}


def _extract_docx_tables(document: Document, *, file_name: str, warnings: list[str]) -> list[dict[str, Any]]:
    tables: list[dict[str, Any]] = []
    try:
        for table_index, table in enumerate(document.tables):
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            markdown = table_rows_to_markdown(rows)
            plain_text = table_rows_to_plain_text(rows)
            if not markdown and not plain_text:
                continue
            tables.append(
                {
                    "tableId": f"{file_name}_table_{table_index}",
                    "fileName": file_name,
                    "pageNumber": None,
                    "tableIndex": table_index,
                    "rowCount": len(rows),
                    "columnCount": max((len(row) for row in rows), default=0),
                    "markdown": markdown,
                    "text": plain_text,
                    "source": "docx-table",
                }
            )
    except Exception as exc:
        warnings.append(f"DOCX table extraction skipped: {exc}")
    return tables


def _extract_docx_comments(zip_file: zipfile.ZipFile, *, file_name: str, warnings: list[str]) -> list[dict[str, Any]]:
    if "word/comments.xml" not in zip_file.namelist():
        return []

    comments: list[dict[str, Any]] = []
    try:
        root = ET.fromstring(zip_file.read("word/comments.xml"))
        for comment_index, comment in enumerate(root.findall(".//w:comment", WORD_NS)):
            text = " ".join(
                node.text.strip()
                for node in comment.findall(".//w:t", WORD_NS)
                if node.text and node.text.strip()
            )
            author = comment.attrib.get(f"{{{WORD_NS['w']}}}author", "")
            if not text:
                continue
            comments.append(
                {
                    "annotationId": f"{file_name}_comment_{comment_index}",
                    "fileName": file_name,
                    "pageNumber": None,
                    "type": "docx-comment",
                    "title": author,
                    "subject": "comment",
                    "content": text,
                    "source": "docx-comment",
                }
            )
    except Exception as exc:
        warnings.append(f"DOCX comment extraction skipped: {exc}")
    return comments


def _extract_docx_notes(zip_file: zipfile.ZipFile, *, file_name: str, note_type: str, warnings: list[str]) -> list[dict[str, Any]]:
    xml_name = f"word/{note_type}.xml"
    if xml_name not in zip_file.namelist():
        return []

    notes: list[dict[str, Any]] = []
    try:
        root = ET.fromstring(zip_file.read(xml_name))
        for note_index, note in enumerate(root):
            text = " ".join(
                node.text.strip()
                for node in note.findall(".//w:t", WORD_NS)
                if node.text and node.text.strip()
            )
            if not text:
                continue
            notes.append(
                {
                    "annotationId": f"{file_name}_{note_type}_{note_index}",
                    "fileName": file_name,
                    "pageNumber": None,
                    "type": note_type,
                    "title": note_type,
                    "subject": note_type,
                    "content": text,
                    "source": note_type,
                }
            )
    except Exception as exc:
        warnings.append(f"DOCX {note_type} extraction skipped: {exc}")
    return notes


def _extract_docx_links(document: Document, *, file_name: str, warnings: list[str]) -> list[dict[str, Any]]:
    links: list[dict[str, Any]] = []
    try:
        for link_index, rel in enumerate(document.part.rels.values()):
            if "hyperlink" not in rel.reltype:
                continue
            links.append(
                {
                    "linkId": f"{file_name}_link_{link_index}",
                    "fileName": file_name,
                    "pageNumber": None,
                    "uri": rel.target_ref,
                    "targetPage": None,
                    "kind": "external",
                    "source": "docx-link",
                }
            )
    except Exception as exc:
        warnings.append(f"DOCX link extraction skipped: {exc}")
    return links


def _build_docx_render_candidates(
    *,
    file_name: str,
    doc_type: str,
    tables: list[dict[str, Any]],
    annotations: list[dict[str, Any]],
    links: list[dict[str, Any]],
    image_count: int,
) -> list[dict[str, Any]]:
    reasons: list[str] = []
    if tables:
        reasons.append("table_shapes")
    if annotations:
        reasons.append("annotations")
    if links:
        reasons.append("document_links")
    if image_count:
        reasons.append("embedded_image")

    meaningful_reasons = [reason for reason in reasons if reason != "document_links"]
    if not meaningful_reasons:
        return []

    return [
        annotate_visual_candidate(
            {
                "fileName": f"{file_name}_render_candidate_document.png",
                "imageId": f"{file_name}_render_candidate_document",
                "docType": f"{doc_type}-IMAGE" if doc_type != "UNKNOWN" else "IMAGE",
                "imageSource": "rendered-page-candidate",
                "pageNumber": None,
                "visualReason": reasons,
                "mimeType": "image/png",
            }
        )
    ]


def extract_docx_document(*, content: bytes, context: dict[str, Any]) -> dict[str, Any]:
    file_name = context["fileName"]
    doc_type = context["docType"]
    extraction_config = context.get("extractionConfig", {})
    extract_tables = bool(extraction_config.get("extractTables", True))
    extract_annotations = bool(extraction_config.get("extractAnnotations", True))
    extract_links = bool(extraction_config.get("extractLinks", True))
    detect_rendered_pages = bool(extraction_config.get("detectRenderedPages", False))
    render_pages = bool(extraction_config.get("renderPages", False))

    document = Document(io.BytesIO(content))
    raw_text_parts = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text and paragraph.text.strip()
    ]

    images: list[dict[str, Any]] = []
    warnings: list[str] = []
    tables = _extract_docx_tables(document, file_name=file_name, warnings=warnings) if extract_tables else []
    links = _extract_docx_links(document, file_name=file_name, warnings=warnings) if extract_links else []
    annotations: list[dict[str, Any]] = []
    if extract_annotations:
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as zip_file:
                annotations.extend(_extract_docx_comments(zip_file, file_name=file_name, warnings=warnings))
                annotations.extend(_extract_docx_notes(zip_file, file_name=file_name, note_type="footnotes", warnings=warnings))
                annotations.extend(_extract_docx_notes(zip_file, file_name=file_name, note_type="endnotes", warnings=warnings))
        except Exception as exc:
            warnings.append(f"DOCX annotation extraction skipped: {exc}")

    image_index = 0
    for rel in document.part.rels.values():
        if "image" not in rel.target_ref:
            continue

        image_bytes = rel.target_part.blob
        extension = rel.target_ref.rsplit(".", 1)[-1].lower()
        images.append(
            {
                "fileName": f"{file_name}_image_{image_index}.{extension}",
                "imageId": f"{file_name}_img_{image_index}",
                "base64": base64.b64encode(image_bytes).decode(),
                "docType": f"{doc_type}-IMAGE" if doc_type != "UNKNOWN" else "IMAGE",
                "imageSource": "embedded-image",
                "pageNumber": None,
                "visualReason": ["embedded_image"],
                "mimeType": f"image/{extension}",
            }
        )
        image_index += 1

    visual_candidates = (
        _build_docx_render_candidates(
            file_name=file_name,
            doc_type=doc_type,
            tables=tables,
            annotations=annotations,
            links=links,
            image_count=len(images),
        )
        if detect_rendered_pages
        else []
    )

    return {
        "fileType": "docx",
        "pageCount": 0,
        "rawText": "\n".join(raw_text_parts),
        "images": images,
        "renderedPages": [],
        "visualCandidates": visual_candidates,
        "tables": tables,
        "annotations": annotations,
        "links": links,
        "warnings": warnings,
        "extractionStats": {
            "compatibilityMode": False,
            "safeTextRichMode": True,
            "renderingEnabled": False,
            "renderPagesRequested": render_pages,
            "renderDetectionEnabled": detect_rendered_pages,
            "extractTablesEnabled": extract_tables,
            "extractAnnotationsEnabled": extract_annotations,
            "extractLinksEnabled": extract_links,
            "paragraphsExtracted": len(raw_text_parts),
            "embeddedImagesExtracted": len(images),
            "renderedPagesGenerated": 0,
            "visualCandidatesDetected": len(visual_candidates),
            "visualCandidatePages": "",
            "tablesExtracted": len(tables),
            "annotationsExtracted": len(annotations),
            "linksExtracted": len(links),
        },
    }
